"""DOCX document extractor."""

from io import BytesIO

from adaptive_rag.core.exceptions import IngestionError


def extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file.

    Args:
        file_bytes: DOCX file content as bytes.

    Returns:
        Extracted text.

    Raises:
        IngestionError: If extraction fails.
    """
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    texts.append(" | ".join(row_text))
        return "\n\n".join(texts)
    except Exception as e:
        raise IngestionError(f"Failed to extract DOCX: {e}") from e
